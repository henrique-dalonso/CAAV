from functools import lru_cache
from pathlib import Path

import docx


CONFIG_DIR = Path(__file__).resolve().parents[1] / "config"
CAMINHO_PROMPT_MESTRE = CONFIG_DIR / "Prompt PUBLICAÇÕES.docx"
CAMINHO_MANUAL_COMPLEMENTAR = CONFIG_DIR / "COMPLEMENTAR - MANUAL.docx"


def _docx_para_texto(caminho):
    documento = docx.Document(caminho)
    partes = [p.text for p in documento.paragraphs]

    for tabela in documento.tables:
        for linha in tabela.rows:
            partes.append(" | ".join(celula.text.strip() for celula in linha.cells))

    return "\n".join(partes)


@lru_cache(maxsize=1)
def carregar_instrucoes_publicacoes():
    """Concatena o prompt mestre + o manual operacional complementar — os
    2 docx guardados em config/ são pensados pra funcionar em conjunto (o
    próprio manual diz isso na seção 17: "usado... sempre em conjunto").

    Diferente do prompt do Extratus (`instrucoes_relatorio.txt`, editável
    e com histórico de versões via prompt_manager de lá), isso aqui ainda
    não tem tela de edição nem versionamento — são os 2 .docx originais,
    lidos e cacheados em memória (só recarrega se o processo reiniciar).
    Versionamento fica pra quando/se for realmente pedido; por ora, trocar
    o conteúdo é editar os arquivos em config/ e reiniciar o servidor.
    """
    prompt_mestre = _docx_para_texto(CAMINHO_PROMPT_MESTRE)
    manual = _docx_para_texto(CAMINHO_MANUAL_COMPLEMENTAR)
    return prompt_mestre + "\n\n" + manual
