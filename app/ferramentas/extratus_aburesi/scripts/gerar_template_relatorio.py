"""
Gera o template Word (config/relatorio_template.docx) usado pelo Extratus.

Segue a especificacao de formatacao do Max (config/instrucoes_relatorio.txt,
PARTE 2): A4, margens 2,5cm, Arial 11, rotulos em negrito, linha divisoria
fina abaixo dos titulos de secao, espacamentos especificos.

Este script so precisa ser rodado de novo se o template for perdido/corrompido
e precisar ser reconstruido do zero. Ajustes normais de visual devem ser
feitos abrindo o .docx gerado direto no Word.

Nao cobre (fora do escopo do MVP, o pipeline hoje trata 1 PDF = 1 processo):
- Secao para incidentes separados (embargos, agravo, reconvencao)
- Separador duplo azul entre multiplos processos no mesmo arquivo
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[4]))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.plataforma.paths import PROJECT_ROOT


CAMINHO_SAIDA = str(
    PROJECT_ROOT / "app" / "ferramentas" / "extratus_aburesi" / "config" / "relatorio_template.docx"
)

CINZA = RGBColor(0x80, 0x80, 0x80)


def adicionar_borda_inferior(paragrafo, cor="808080", tamanho=6):
    p_pr = paragrafo._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(tamanho))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), cor)
    p_borders.append(bottom)
    p_pr.append(p_borders)


def configurar_pagina(documento):
    secao = documento.sections[0]
    secao.page_height = Cm(29.7)
    secao.page_width = Cm(21.0)
    secao.top_margin = Cm(2.5)
    secao.bottom_margin = Cm(2.5)
    secao.left_margin = Cm(2.5)
    secao.right_margin = Cm(2.5)


def configurar_fonte_padrao(documento):
    estilo = documento.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(11)
    estilo.paragraph_format.space_after = Pt(0)
    estilo.paragraph_format.line_spacing = 1.0


def campo_rotulo(documento, rotulo, variavel, espaco_depois=4):
    paragrafo = documento.add_paragraph()
    paragrafo.paragraph_format.space_after = Pt(espaco_depois)
    paragrafo.paragraph_format.line_spacing = 1.0

    run_rotulo = paragrafo.add_run(f"{rotulo}: ")
    run_rotulo.bold = True

    paragrafo.add_run("{{ " + variavel + " }}")

    return paragrafo


def titulo_secao(documento, texto):
    paragrafo = documento.add_paragraph()
    paragrafo.paragraph_format.space_before = Pt(12)
    paragrafo.paragraph_format.space_after = Pt(6)

    run = paragrafo.add_run(texto)
    run.bold = True
    run.font.size = Pt(11)

    adicionar_borda_inferior(paragrafo)

    return paragrafo


def construir_template():
    documento = Document()

    configurar_pagina(documento)
    configurar_fonte_padrao(documento)

    # PARTE 1 - cabecalho de identificacao
    titulo = documento.add_paragraph()
    run_titulo = titulo.add_run("RELATÓRIO PROCESSUAL")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    titulo.paragraph_format.space_after = Pt(12)

    campo_rotulo(documento, "TIPO DA AÇÃO", "tipo_acao")
    campo_rotulo(documento, "N° PROCESSO", "numero_processo")
    campo_rotulo(documento, "INCIDENTE", "incidente")
    campo_rotulo(documento, "VALOR DA CAUSA", "valor_causa")
    campo_rotulo(documento, "VALOR DA DÍVIDA AJUIZADA", "valor_divida")
    campo_rotulo(documento, "AUTOR", "autor")
    campo_rotulo(documento, "RÉU", "reu")
    campo_rotulo(documento, "BEM", "bem")
    campo_rotulo(documento, "CONTRATO", "contrato")
    campo_rotulo(documento, "COMARCA/TRIBUNAL", "comarca")

    # Cronologia processual
    titulo_secao(documento, "CRONOLOGIA PROCESSUAL")

    documento.add_paragraph("{% for evento in cronologia %}")

    paragrafo_evento = documento.add_paragraph()
    paragrafo_evento.paragraph_format.space_after = Pt(6)
    run_data = paragrafo_evento.add_run("{{ evento.data }}")
    run_data.bold = True
    paragrafo_evento.add_run(" – {{ evento.ator }} – {{ evento.descricao }}")

    documento.add_paragraph("{% endfor %}")

    # Parecer do escritorio
    titulo_secao(documento, "PARECER DO ESCRITÓRIO")

    paragrafo_parecer = documento.add_paragraph("{{ parecer }}")
    paragrafo_parecer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo_parecer.paragraph_format.space_after = Pt(8)

    campo_rotulo(documento, "Data publicação/ciência", "data_publicacao")
    campo_rotulo(documento, "Prazo Fatal ED", "prazo_fatal_ed")
    campo_rotulo(documento, "Prazo Fatal", "prazo_fatal")
    campo_rotulo(documento, "Status atual", "status_atual")

    documento.save(CAMINHO_SAIDA)
    print(f"Template salvo em: {CAMINHO_SAIDA}")


if __name__ == "__main__":
    construir_template()
